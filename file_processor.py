import os
import logging
import fitz  # PyMuPDF
from docx import Document

def process_file(filepath):
    """
    Extract text from PDF or DOCX files
    Returns extracted text content
    """
    try:
        file_extension = os.path.splitext(filepath)[1].lower()
        
        if file_extension == '.pdf':
            return extract_pdf_text(filepath)
        elif file_extension == '.docx':
            return extract_docx_text(filepath)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        logging.error(f"Error processing file {filepath}: {str(e)}")
        raise

def extract_pdf_text(filepath):
    """Extract text from PDF using PyMuPDF"""
    try:
        doc = fitz.open(filepath)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        
        doc.close()
        return text.strip()
        
    except Exception as e:
        logging.error(f"Error extracting PDF text: {str(e)}")
        raise

def extract_docx_text(filepath):
    """Extract text from DOCX using python-docx"""
    try:
        doc = Document(filepath)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
        
    except Exception as e:
        logging.error(f"Error extracting DOCX text: {str(e)}")
        raise

def validate_file_content(text):
    """Validate that extracted text looks like a resume"""
    if not text or len(text.strip()) < 50:
        return False, "File appears to be empty or too short"
    
    # Basic checks for resume content
    common_resume_words = ['experience', 'education', 'skills', 'work', 'employment', 'degree', 'university', 'college']
    text_lower = text.lower()
    
    found_words = sum(1 for word in common_resume_words if word in text_lower)
    
    if found_words < 2:
        return False, "File does not appear to contain resume content"
    
    return True, "Valid resume content"
